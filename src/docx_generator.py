# src/docx_generator.py
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from PIL import Image

def clean_text(text):
    """Simply clean up the text and retain readable characters (though Word supports Unicode, 
    so there's no need to clear special characters)"""
    
    return re.sub(r'\n{3,}', '\n\n', text).strip()

def generate_docx_report(symbol, company_name, insights, suggestions,
                        md_report_path, figure_path, output_path="reports/final_report.docx"):
    """
    The parameters remain the same as before, 
    but it is recommended to change the output extension to docx
    """
    # ---- read Markdown detail（if have） ----
    try:
        with open(md_report_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except:
        md_content = f"# {company_name} Analysis Report"

    # peel Markdown sign
    plain_text = re.sub(r'[#*`>_\-]', '', md_content)
    plain_text = re.sub(r'\n{3,}', '\n\n', plain_text)

    # ---- create Word docx ----
    doc = Document()

    # Set page margins (optional)
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ---- title ----
    title_para = doc.add_heading(f"{company_name} ({symbol}) – Sentiment Analysis Report", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # date
    date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # space

    # ---- AI Insights ----
    doc.add_heading("AI Key Insights", level=2)
    for ins in insights:
        # Add bullet points (built-in in Word)
        p = doc.add_paragraph(ins, style='ListBullet')
    doc.add_paragraph()  # sapce

    # ---- Recommendations ----
    doc.add_heading("Recommendations", level=2)
    for sug in suggestions:
        p = doc.add_paragraph(sug, style='ListBullet')
    doc.add_paragraph()

    # ---- Charts ----
    if os.path.exists(figure_path):
        doc.add_heading("Charts", level=2)
        try:
            # Obtain the image size to control the display width 
            # (limit the maximum width to 6 inches)
            with Image.open(figure_path) as img:
                w, h = img.size
                # Convert pixels to inches 
                # (assuming 96 dpi, can be adjusted appropriately)
                width_in_inches = min(w / 96, 6.0)  # The maximum width is 6 inches
                # If the aspect ratio of the image is abnormal, it can be adjusted
                doc.add_picture(figure_path, width=Inches(width_in_inches))
        except:
            doc.add_picture(figure_path, width=Inches(6.0))
        doc.add_paragraph()  # space
    else:
        doc.add_paragraph("(Chart image not found)")

    # ---- Detailed Data ----
    doc.add_heading("Detailed Data", level=2)
    # Add paragraph by paragraph, separated by line breaks
    for line in plain_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())

    # ---- declaration ----
    
    doc.add_paragraph()  # space
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Disclaimer: For educational purposes only. Not financial advice.")
    run.italic = True
    run.font.size = Pt(9)

    # second date
    date_footer = doc.add_paragraph()
    date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = date_footer.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run2.italic = True
    run2.font.size = Pt(9)

    # ---- save ----
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Word report generated: {output_path}")
    return output_path
